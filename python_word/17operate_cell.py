from docx import Document
from docx.shared import Mm

doc = Document()
table = doc.add_table(5, 3)

for row in table.rows[2:]:
    for j, cell in enumerate(row.cells):
        print(f"cell {row._index, j}: ", cell)

for cell in table.row_cells(2):
    print("row cell: ", cell)

for cell in table.column_cells(2):
    print("column cell: ", cell)

cell_1 = table.cell(1, 1)

# 修改单元格
cell_1.text = "冰冷的希望"
print(cell_1.text)

paragraph = cell_1.paragraphs[0]
paragraph.add_run("!")
paragraph.add_run("冰冰好帅")
cell_2 = table.cell(3, 2)
cell_2.add_paragraph("大数据协会")
cell_2.add_paragraph().add_run().add_picture("./input/luffy.png", width=Mm(20))

# 单元格合并
cell_1.merge(cell_2)

doc.save('./output/test.docx')