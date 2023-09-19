from docx import Document

doc_file = Document('./output/test.docx')

p = doc_file.paragraphs[0]

# run = p.add_run("aaa")
# run2 = p.add_run("bbb")

print(len(p.runs))

for r in p.runs:
    if r.text == 'aaa':
        r.text = 'ccc'
        r.clear()
    print(r.text)


doc_file.save('./output/test.docx')