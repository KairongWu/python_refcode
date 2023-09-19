from docx import Document
from docx.shared import Cm

doc = Document()

pic = doc.add_picture("./input/luffy.png", width=Cm(5))
# doc.add_picture("./input/luffy.png")

doc.paragraphs[0].runs[0].add_picture("./input/luffy.png", width=Cm(5))

print(type(pic), pic.type)

print(doc.inline_shapes)
print(len(doc.inline_shapes))

# 行内元素: 可以把多个元素放在同一行显示，inline_shapes，runs
# 快捷元素：不能放在同一行，paragraphs
for shape in doc.inline_shapes:
    print("type: ", shape.type)

doc.save('./output/test.docx')