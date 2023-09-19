from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# 社团证明案例


college = '化学化工'
major = '化学'
class_ = "1601"
name = '冰冰'
join_year = '2021-2022'
association = '大数据协会'
prove_date = '2022年12月30日'

doc_file = Document()

run_list = list()  # 保存需要设置下户线的 Run 对象

# 标题段落
title = doc_file.add_paragraph()  # type: Paragraph
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中
title_run = title.add_run("证明")  # type: Run
title_run.font.size = Pt(22)  # 设置字号为二号， 即 22 磅
title_run.font.name = "Time New Roman"  # 要先设置西文才能设置中文
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 正文段落
body = doc_file.add_paragraph()  # type: Paragraph
body.add_run("兹证明，")
run_list.append(body.add_run(f"   {college}   "))
body.add_run("学院")
run_list.append(body.add_run(f"   {major}   "))
body.add_run("系")
run_list.append(body.add_run(f"   {class_}   "))
body.add_run("班")
run_list.append(body.add_run(f"   {name}   "))
body.add_run(f"同学于{join_year}学年加入")
run_list.append(body.add_run(f"   {association}   "))
body.add_run(
    "（社团名称），鉴于该同学在我社期间各方面表现优秀，特此证明，"
    "并建议按照综测细则给予该同学相应综合量化成绩加分。"
)
for run in run_list:
    run.font.underline = True  # 遍历添加下划线

font_size = 14  # 正文四号字体，即 14 磅
body.style.font.size = Pt(font_size)  # 设置段落字体为 14 磅
body.paragraph_format.first_line_indent = Pt(font_size * 2)  #
body.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5 倍
body.style.font.name = "Time New Roman"  # 设置西文为新罗马
body.style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置

# 署名和日期
p_association = doc_file.add_paragraph(f"{association}")
p_association.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p_prove_date = doc_file.add_paragraph(f"{prove_date}")
p_prove_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc_file.save('./output/test.docx')