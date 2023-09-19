from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

doc = Document()
section1 = doc.sections[0]
doc.add_paragraph("第一节的段落")

header = section1.header  # 此处没有 s 因为一个 section只有一个页眉
footer = section1.footer
print(header)
print(footer)

header.paragraphs[0].add_run('这是页眉1')
footer.paragraphs[0].add_run('这是页脚1')

section1.header_distance = Cm(1.0)
section1.footer_distance = Cm(1.2)
section1.header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

section2 = doc.add_section()
doc.add_paragraph("第二节的段落")
section2.header.is_linked_to_previous = False
section2.header.paragraphs[0].text = "这是页眉2"
section2.footer.is_linked_to_previous = False
section2.footer.paragraphs[0].text = "这是页脚2"

doc.save('./output/test.docx')