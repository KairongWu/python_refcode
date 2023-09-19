from docx import Document

doc = Document()

doc.add_heading("标题 1", level=1)
doc.add_heading("标题 2", level=2)
doc.add_heading("标题 3", level=3)
doc.add_heading("标题 4", level=4)
doc.add_heading("标题 5", level=5)

doc.save("./output/test.docx")