from docx import Document
from docx.shared import Pt, Mm
from docx.text.run import Run

# 1 Emu = 1
# I Twips = 635 Emu
# 1 Pt = 12700 Emu
# 1 Mm = 36000 Emu
# 1 Cm = 360000 Emu
# 1 Inches = 914400 Emu

doc = Document()
run = doc.add_paragraph().add_run("冰冷的希望")  # type: Run

run.font.size = Pt(14)

p = Pt(14)
print(p.emu)
print(p.mm)
print(Mm(20).inches)