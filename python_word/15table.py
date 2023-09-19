from docx import Document

doc = Document()

table = doc.add_table(3, 2)
print(type(table))

print(doc.tables)
for t in doc.tables:
    print(t)

doc.save('./output/test.docx')