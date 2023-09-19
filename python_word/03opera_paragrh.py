from docx import Document

docx_file = Document()


##### one paragraph test
# print(type(docx_file))
# print(len(docx_file.paragraphs))
#
# paragrh_1 = docx_file.add_paragraph(text="万岁，加油")
# print(type(paragrh_1))
# print(len(docx_file.paragraphs))
#
# paragrh_1.text = '1111'
# # paragrh_1.text = ''
# paragrh_1.clear()
# print(paragrh_1.text)


#### multi paragraph test
paragrh_1 = docx_file.add_paragraph(text="万岁，加油1")
docx_file.add_page_break()   # 换页符
paragrh_2 = docx_file.add_paragraph(text="万岁，加油2")
paragrh_3 = docx_file.add_paragraph(text="万岁，加油3")


paragrh_2.insert_paragraph_before("这是插进来的段落")

print(len(docx_file.paragraphs))


docx_file.save("./output/test.docx")