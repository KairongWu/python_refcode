from docx import Document
# 奇偶页与首页

doc = Document()

# doc.settings.odd_and_even_pages_header_footer = True
#
section = doc.sections[0]
# section.header.paragraphs[0].text = "奇数页页眉"
# section.footer.paragraphs[0].text = "奇数页页脚"
#
# even_page_header = section.even_page_header
# even_page_footer = section.even_page_footer
# print(type(even_page_header))
# print(type(even_page_footer))
# even_page_header.paragraphs[0].text = "偶数页页眉"
# even_page_footer.paragraphs[0].text = "偶数页页脚"


# 首页与其他页不同
section.different_first_page_header_footer = True
section.first_page_header.paragraphs[0].text = "首页页眉"
section.first_page_footer.paragraphs[0].text = "首页页脚"

section.header.paragraphs[0].text = "其他页页眉"
section.footer.paragraphs[0].text = "其他页页脚"

doc.add_page_break()
doc.add_page_break()
doc.add_page_break()


doc.save('./output/test.docx')