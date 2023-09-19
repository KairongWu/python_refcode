from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

pic = doc.add_picture('./input/luffy.png')
# pic = doc.add_picture('./input/luffy.png', Mm(50))
# pic = doc.add_picture('./input/luffy.png', height=Mm(50), width=Mm(80))

print(pic.height)
print(pic.width.mm)

# pic.width = Mm(60)
# print(pic.height)
# print(pic.width.mm)

pic.width = int(pic.width * 0.1)
pic.height = int(pic.height * 0.1)

# doc.paragraphs[0].aligment = WD_PARAGRAPH_ALIGNMENT.LEFT
doc.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# doc.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
# doc.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE

doc.save('./output/test.docx')