from docx import Document
from docx.enum.text import WD_UNDERLINE, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

doc_file = Document()

paragraph1 = doc_file.add_paragraph("这是第一个段落")  # type: Paragraph
run1 = paragraph1.add_run("追加的文字")
run1.bold = True  # 加粗
run1.italic = True  # 斜体
run1.underline = True  # 下划线
# run1.text = "修改之后的文字"  # 修改文本
run1.font.name = "Times New Roman"  # 设置西文是新罗马字体
run1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置中文是宋体
run1.font.size = Pt(30)  # 字号大小
run1.font.bold = False  # 是否加粗
run1.font.italic = False  # 是否斜体
run1.font.underline = False  # 是否下划线
# run1.font.underline = WD_UNDERLINE.DOUBLE  # 设置为双下划线
# 查看所有下划线类型
for line_type in WD_UNDERLINE.__members__:
    print(line_type.name)
run1.font.shadow = True  # 是否阴影
run1.font.strike = True  # 是否删除线
# run1.font.double_strike = True  # 是否双删除线
run1.font.color.rgb = RGBColor(56, 36, 255)  # 字体颜色
# run1.font.color.rgb = RGBColor.from_string("ff056")  # 字体颜色
run1.font.highlight_color = WD_COLOR_INDEX.YELLOW  # 文本高亮颜色， 此次设置为黄色
# 查看所有支持的高亮颜色
for color in WD_COLOR_INDEX.__members__:
    print(color.name)

doc_file.save("./output/test.docx")