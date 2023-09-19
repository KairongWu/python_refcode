from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt
from docx.text.paragraph import Paragraph

doc_file = Document()

# p1 = doc_file.add_paragraph("左对齐")
# p2 = doc_file.add_paragraph("右对齐")
# p3 = doc_file.add_paragraph("居中对齐")
# p4 = doc_file.add_paragraph("两端对齐")
# p5 = doc_file.add_paragraph("分散对齐")
#
# print(len(doc_file.paragraphs))
#
# p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
# p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# p4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
# p5.alignment = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE




# styles = doc_file.styles
# print(len(styles))
# paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]   # alt + enter 可快速导入确实模块
# print(len(paragraph_styles))
# for style in paragraph_styles:
#     print(style.name)



# paragraph1 = doc_file.add_paragraph("这是第一个段落")  # type: Paragraph
# paragraph1.style = "List Bullet"  # 设置段落样式为无序 1
#
# # 有序段落央视
# doc_file.add_paragraph("测试段落 List Number", style="List Number")
# doc_file.add_paragraph("测试段落 List Number", style="List Number 2")
# doc_file.add_paragraph("测试段落 List Number", style="List Number 3")
#
# # 无序段落样式2
# doc_file.add_paragraph("测试段落 List Bullet 2", style="List Bullet")
# doc_file.add_paragraph("测试段落 List Bullet 2", style="List Bullet 2")
# doc_file.add_paragraph("测试段落 List Bullet 2", style="List Bullet 3")



# paragraph1 = doc_file.add_paragraph("这是一个普通的段落这是一个普通的段落这是一个普通的段落")
# # paragraph1.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 行距固定值
# # paragraph1.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # 多倍行距
# # paragraph1.paragraph_format.line_spacing = 1.5  # 行间距， 1.5 倍行距
# paragraph1.paragraph_format.line_spacing = Pt(20)  # 行间距，固定值 20 磅
# paragraph1.paragraph_format.first_line_indent = Pt(20)  # 首行缩进 10 磅
# paragraph1.paragraph_format.space_before = Pt(30)  # 段前 30 磅
# paragraph1.paragraph_format.space_after = Pt(15)  # 段后 15 磅



p = doc_file.add_paragraph("测试段落 List Bullet 2", style="List Bullet")
print(p.style.name)
p.style.delete()
print(p.style.name)


doc_file.save('./output/test.docx')